from docx import Document


def markdown_to_docx(text):
    out = Document()

    for i in text.split('\n')[1::]:
        if '#' in i.split()[0]:
            num = i.count('#')
            if i.endswith('#'):
                out.add_heading(i[num // 2:-1 - (num // 2)], level=num // 2)
            else:
                out.add_heading(i[num::], level=num)

        elif i.split()[0] == '*' or i.split() or '+' in i.split():
            out.add_paragraph(i[2::], style='List Bullet')
        elif type(i.split()[:-1]) == int:
            out.add_paragraph(' '.join(i.split()[1::]), style='List Number')
        elif not i:
            out.add_paragraph('')
        elif (i.startswith('_') and i.endwith('_')) or (i.startswith('*') and i.endswith('*')):
            if i.startswith('_'):
                counter = i.count('_') // 2
            else:
                counter = i.count('*') // 2
            if counter == 1:
                out.add_paragraph(i[1:-1]).italic = True
            elif counter == 2:
                out.add_paragraph(i[2:-2]).strong = True
            else:
                out.add_paragraph(i[3:-3]).strong, italic = True

    out.save('res.docx')
